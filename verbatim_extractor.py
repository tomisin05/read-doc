#!/usr/bin/env python3
"""
Verbatim Debate Card Extractor
================================
Extracts highlighted and/or underlined text from a verbatim-formatted .docx file.
Keeps the full card structure (tags, cites, block headers) but removes un-marked body text.

Usage:
    python verbatim_extractor.py input.docx [-H] [-U]

Options:
    -H    Extract only highlighted text
    -U    Extract only underlined text
    (default: extract both highlighted and underlined text)
"""

import argparse
import sys
import copy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX
from lxml import etree


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def run_is_highlighted(run):
    """Return True if this run has ANY highlight color set (any color, not 'none')."""
    rpr = run._r.find(qn('w:rPr'))
    if rpr is None:
        return False
    
    # Check standard highlight
    highlight = rpr.find(qn('w:highlight'))
    if highlight is not None:
        val = highlight.get(qn('w:val'), '')
        if val.lower() not in ('', 'none'):
            return True
    
    # Check background shading (often used when pasting from web)
    shd = rpr.find(qn('w:shd'))
    if shd is not None:
        fill = shd.get(qn('w:fill'), '')
        # 'auto' or empty means no shading
        if fill.lower() not in ('', 'auto', 'ffffff', 'none'):
            return True
    
    return False


def run_is_underlined(run):
    """Return True if this run has underline formatting."""
    rpr = run._r.find(qn('w:rPr'))
    if rpr is None:
        return False
    u = rpr.find(qn('w:u'))
    if u is None:
        return False
    val = u.get(qn('w:val'), 'single')  # default to 'single' if no value
    # 'none' explicitly turns underline off
    return val.lower() != 'none'


def run_is_bold(run):
    """Return True if this run is bold (used to detect tags/cites)."""
    rpr = run._r.find(qn('w:rPr'))
    if rpr is None:
        return False
    b = rpr.find(qn('w:b'))
    return b is not None


def paragraph_is_structural(para):
    """
    Heuristic: a paragraph is a 'structural' paragraph (tag, cite, block header)
    if ALL of its non-empty runs are bold, or if the paragraph style suggests a heading.
    These should be kept regardless of highlight/underline status.
    """
    style_name = (para.style.name or '').lower()
    if any(k in style_name for k in ('heading', 'block', 'tag', 'cite', 'title')):
        return True

    runs = [r for r in para.runs if r.text.strip()]
    if not runs:
        return False  # empty paragraph

    # If every non-empty run is bold → treat as structural
    return all(run_is_bold(r) for r in runs)


def paragraph_has_marked_runs(para, mode):
    """Return True if any run in this paragraph passes the filter."""
    for run in para.runs:
        if not run.text.strip():
            continue
        if mode == 'highlighted' and run_is_highlighted(run):
            return True
        if mode == 'underlined' and run_is_underlined(run):
            return True
        if mode == 'both' and (run_is_highlighted(run) or run_is_underlined(run)):
            return True
    return False


def filter_paragraph_runs(para, mode):
    """
    Remove runs that don't pass the filter. Returns True if any runs remain.
    Structural paragraphs (tags/cites) are returned as-is with all runs intact.
    """
    if paragraph_is_structural(para):
        return True  # keep everything

    runs_to_remove = []
    prev_kept = False
    for run in para.runs:
        text = run.text.strip()
        if not text:
            continue
        keep = False
        if mode == 'highlighted':
            keep = run_is_highlighted(run)
        elif mode == 'underlined':
            keep = run_is_underlined(run)
        else:  # both
            keep = run_is_highlighted(run) or run_is_underlined(run)

        if not keep:
            runs_to_remove.append(run)
        else:
            # Add space before this run if previous run was also kept
            if prev_kept and not run.text.startswith(' '):
                run.text = ' ' + run.text
            prev_kept = True

    for run in runs_to_remove:
        run._r.getparent().remove(run._r)

    # Check if anything remains
    remaining = [r for r in para.runs if r.text.strip()]
    return len(remaining) > 0


# ---------------------------------------------------------------------------
# Main extraction logic
# ---------------------------------------------------------------------------

def extract(input_path, output_path, mode='both'):
    print(f"[•] Loading: {input_path}")
    doc = Document(input_path)
    out_doc = Document(input_path)  # start from same doc to preserve styles/formatting

    # We'll rebuild the body by removing paragraphs that have nothing left
    # Work on out_doc directly

    body = out_doc.element.body
    paragraphs = out_doc.paragraphs

    paras_to_remove = []
    prev_was_structural = False

    for para in paragraphs:
        is_structural = paragraph_is_structural(para)
        
        # Keep structural paragraphs and the paragraph immediately after (cite)
        if is_structural or prev_was_structural:
            prev_was_structural = is_structural
            continue
        
        prev_was_structural = False

        # Check if there are any marked runs
        if not paragraph_has_marked_runs(para, mode):
            # No marked content at all — queue for removal
            paras_to_remove.append(para)
        else:
            # Has some marked runs — strip unmarked ones
            has_content = filter_paragraph_runs(para, mode)
            if not has_content:
                paras_to_remove.append(para)

    # Remove empty paragraphs
    removed = 0
    for para in paras_to_remove:
        p = para._element
        p.getparent().remove(p)
        removed += 1

    # Clean up consecutive empty paragraphs that might be left over
    all_paras = out_doc.element.body.findall('.//' + qn('w:p'))
    prev_empty = False
    for p_el in all_paras:
        texts = ''.join(t.text or '' for t in p_el.iter(qn('w:t')))
        is_empty = not texts.strip()
        if is_empty and prev_empty:
            p_el.getparent().remove(p_el)
        prev_empty = is_empty

    print(f"[✓] Removed {removed} body paragraphs with no marked content")
    out_doc.save(output_path)
    print(f"[✓] Saved to: {output_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description='Verbatim Debate Card Extractor — keeps highlighted/underlined runs only'
    )
    parser.add_argument('input', help='Input .docx file')
    parser.add_argument('-H', action='store_true', help='Extract only highlighted text')
    parser.add_argument('-U', action='store_true', help='Extract only underlined text')
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"[✗] File not found: {input_path}")
        sys.exit(1)

    if args.H:
        mode = 'highlighted'
    elif args.U:
        mode = 'underlined'
    else:
        mode = 'both'

    output_path = input_path.parent / f"{input_path.stem}_read-doc{input_path.suffix}"
    extract(str(input_path), str(output_path), mode=mode)


if __name__ == '__main__':
    main()
